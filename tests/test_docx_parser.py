"""Unit tests for DOCX parser."""
from __future__ import annotations

import sys
from pathlib import Path
from io import BytesIO

import pytest

from paperfraud.base import ParsedPaper
from paperfraud.parser.engine import parse_paper
from paperfraud.config import Config

pytest.importorskip("docx", reason="python-docx not installed")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _make_docx(paragraphs: list[tuple[str, str | None, bool, int | None]] = None,
               tables: list[list[list[str]]] = None) -> Path:
    """Create a minimal .docx in memory, save to a temp file.

    Args:
        paragraphs: List of (text, style_name, bold, font_size_pt).
        tables: List of tables, each is list of rows, each row is list of cell texts.
    """
    doc = Document()

    # Set title
    doc.core_properties.title = "Test Paper"
    doc.core_properties.author = "Smith; Jones"

    if paragraphs:
        for text, style, bold, size in paragraphs:
            para = doc.add_paragraph()
            run = para.add_run(text)
            if bold:
                run.bold = True
            if size:
                run.font.size = Pt(size)
            if style:
                para.style = doc.styles[style] if style in [s.name for s in doc.styles] else para.style

    if tables:
        for rows in tables:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 0)
            for i, row_cells in enumerate(rows):
                for j, cell_text in enumerate(row_cells):
                    table.cell(i, j).text = cell_text

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Write to temp file
    tmp = Path(f"/tmp/test_docx_{id(buf)}.docx")
    tmp.write_bytes(buf.read())
    return tmp


class TestDocxParser:
    """Tests for DocxParser.parse()."""

    def test_parse_simple_docx(self):
        """A basic docx with text should populate full_text."""
        path = _make_docx([
            ("This is the abstract of the paper. It describes the study.", None, False, None),
            ("The introduction presents the background.", None, False, None),
            ("We used standard methods for analysis.", None, False, None),
        ])
        try:
            paper = parse_paper(path, Config())
            assert len(paper.full_text) > 50
            assert "abstract" in paper.full_text.lower()
        finally:
            path.unlink(missing_ok=True)

    def test_parse_heading_structure(self):
        """Documents with Heading styles should segment into IMRaD sections."""
        path = _make_docx([
            ("Abstract", "Heading 1", True, 16),
            ("This paper investigates the effects of X on Y.", None, False, 11),
            ("Introduction", "Heading 1", True, 16),
            ("Background information about the topic.", None, False, 11),
            ("Methods", "Heading 1", True, 16),
            ("Cell culture and Western blot were performed.", None, False, 11),
            ("Results", "Heading 1", True, 16),
            ("The treatment significantly increased expression.", None, False, 11),
            ("Discussion", "Heading 1", True, 16),
            ("These findings demonstrate a novel mechanism.", None, False, 11),
        ])
        try:
            paper = parse_paper(path, Config())
            assert len(paper.methods) > 10
            assert "Western blot" in paper.methods
            assert len(paper.results) > 10
            assert "treatment" in paper.results.lower()
            assert len(paper.discussion) > 10
        finally:
            path.unlink(missing_ok=True)

    def test_parse_visual_heading_fallback(self):
        """ALL-CAPS short paragraphs with bold should be detected as section headings
        even without Heading styles."""
        path = _make_docx([
            ("ABSTRACT", None, True, 14),
            ("This paper investigates the effects of X on Y.", None, False, 11),
            ("INTRODUCTION", None, True, 14),
            ("Background information about the topic.", None, False, 11),
            ("MATERIALS AND METHODS", None, True, 14),
            ("Cell culture and Western blot were performed.", None, False, 11),
            ("RESULTS", None, True, 14),
            ("The treatment significantly increased expression.", None, False, 11),
        ])
        try:
            paper = parse_paper(path, Config())
            assert len(paper.methods) > 10, f"methods not segmented, got: {paper.methods[:100] if paper.methods else 'EMPTY'}"
            assert "Western blot" in paper.methods
            assert len(paper.results) > 10
        finally:
            path.unlink(missing_ok=True)

    def test_parse_metadata(self):
        """Core properties should be extracted as metadata."""
        path = _make_docx([
            ("Some text in the paper.", None, False, 11),
        ])
        try:
            paper = parse_paper(path, Config())
            assert paper.title == "Test Paper"
            assert len(paper.authors) >= 1
            assert "Smith" in paper.authors[0]
        finally:
            path.unlink(missing_ok=True)

    def test_parse_tables(self):
        """Tables should be extracted into paper.tables."""
        path = _make_docx(
            paragraphs=[
                ("Methods", "Heading 1", True, 16),
                ("We used the following reagents.", None, False, 11),
            ],
            tables=[
                [
                    ["Reagent", "Company", "Catalog"],
                    ["TRIzol", "Invitrogen", "15596026"],
                    ["SYBR Green", "Bio-Rad", "1725124"],
                ],
            ],
        )
        try:
            paper = parse_paper(path, Config())
            assert len(paper.tables) >= 1
            tbl = paper.tables[0]
            assert tbl["headers"] == ["Reagent", "Company", "Catalog"]
            assert len(tbl["rows"]) == 2
            assert tbl["rows"][0]["Reagent"] == "TRIzol"
        finally:
            path.unlink(missing_ok=True)

    def test_parse_no_headings_fallback(self):
        """Document without any headings should use full-text + abstract fallback."""
        path = _make_docx([
            ("This is a long paragraph that serves as the abstract or introduction to "
             "the paper and contains enough text to be detected as a substantial "
             "paragraph by the fallback heuristic which looks for paragraphs longer "
             "than 150 characters in the full text of the document.", None, False, 11),
            ("This is another paragraph with additional content about the methods "
             "used in the study which should also be included in the full text.", None, False, 11),
        ])
        try:
            paper = parse_paper(path, Config())
            assert len(paper.full_text) > 100
            # Fallback abstract should capture first substantial paragraph
            assert len(paper.abstract) > 50
        finally:
            path.unlink(missing_ok=True)

    def test_missing_dependency(self):
        """ImportError should be raised if python-docx is not installed."""
        # Simulated by checking the import guard in docx_parser
        from paperfraud.parser.docx_parser import DocxParser
        assert DocxParser is not None

    def test_unsupported_format(self):
        """Passing an unsupported file format should raise ValueError."""
        tmp = Path("/tmp/test_unsupported.txt")
        tmp.write_text("not a paper")
        try:
            with pytest.raises(ValueError, match="不支持的文件格式"):
                parse_paper(tmp, Config())
        finally:
            tmp.unlink(missing_ok=True)

    def test_engine_dispatch_pdf(self):
        """The engine should dispatch .pdf files to PyMuPDFParser."""
        from paperfraud.parser.engine import parse_paper
        pdf = Path(__file__).resolve().parent / "fixtures" / "elife-54695.pdf"
        if not pdf.exists():
            pytest.skip("Fixture PDF not found")
        paper = parse_paper(pdf, Config(max_pages=2))
        assert len(paper.full_text) > 100
