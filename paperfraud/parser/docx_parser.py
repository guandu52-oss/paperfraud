"""python-docx based .docx parser.

Maps Word document content into the ParsedPaper structure.
Section segmentation uses three-level fallback:
  1. Heading styles (Heading 1/2/3)
  2. Visual heuristics (all-caps short paragraphs, bold + large font >= 14pt)
  3. Full text fallback with first-paragraph abstract heuristic
"""

from __future__ import annotations

import atexit
import re
import shutil
import tempfile
from pathlib import Path

from paperfraud.base import ParsedPaper

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None
    RT = None


class DocxParser:
    """Parse a .docx file into ParsedPaper.

    Strategy:
    - Paragraph text grouped by heading style or visual cues into IMRaD sections
    - Images extracted via relationship XML traversal
    - Tables extracted into paper.tables list of dicts
    """

    # Canonical section keywords, lowercased
    SECTION_KEYWORDS: dict[str, list[str]] = {
        "abstract": ["abstract", "summary", "background"],
        "introduction": ["introduction", "background"],
        "methods": [
            "methods", "materials and methods", "experimental procedures",
            "materials & methods", "methodology", "experimental design",
            "materials", "experimental",
        ],
        "results": ["results", "findings"],
        "discussion": ["discussion", "conclusions", "conclusion",
                        "general discussion", "summary and conclusions"],
    }

    def parse(self, file_path: Path, skip_images: bool = True) -> ParsedPaper:
        if Document is None:
            raise ImportError(
                "Word (.docx) support requires python-docx. "
                "Install with: pip install 'paperfraud-detect[docx]'"
            )

        doc = Document(str(file_path))
        paper = ParsedPaper(file_path=file_path)

        full_text_parts: list[str] = []
        image_dir: Path | None = None

        if not skip_images:
            image_dir = Path(tempfile.mkdtemp(prefix="paperfraud_docx_"))
            paper._tmp_dir = str(image_dir)
            atexit.register(shutil.rmtree, str(image_dir), ignore_errors=True)

        # Collect all paragraphs with their style and formatting info
        paragraphs: list[dict] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs.append({"text": "", "style": None, "bold": False, "size": None})
                continue

            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith("Heading") if style_name else False

            # Get font info from first run
            bold = False
            size = None
            for run in para.runs:
                if run.bold:
                    bold = True
                if run.font.size:
                    size = run.font.size.pt
                    break

            paragraphs.append({
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
                "bold": bold,
                "size": size,
                "heading_level": int(style_name.replace("Heading", "").strip()) if is_heading and style_name != "Heading" else 1,
            })
            full_text_parts.append(text)

        paper.full_text = "\n\n".join(full_text_parts)

        # ── Segment sections ──
        self._segment_sections(paper, paragraphs)

        # ── Extract metadata ──
        self._extract_metadata(paper, doc)

        # ── Extract tables ──
        self._extract_tables(paper, doc)

        # ── Extract images ──
        if image_dir is not None:
            self._extract_images(paper, doc, image_dir)

        return paper

    def _segment_sections(self, paper: ParsedPaper, paragraphs: list[dict]) -> None:
        """Segment text into IMRaD sections using three-level fallback."""
        # Level 1: Try heading styles
        sections = self._segment_by_headings(paragraphs)
        if sections and len(sections) >= 2:
            self._apply_sections(paper, sections)
            if paper.methods or paper.results:
                return

        # Level 2: Visual heuristics for docs without heading styles
        sections = self._segment_by_visual_cues(paragraphs)
        if sections and len(sections) >= 2:
            self._apply_sections(paper, sections)
            if paper.methods or paper.results:
                return

        # Level 3: Full text fallback
        self._fallback_abstract(paper)

    def _segment_by_headings(self, paragraphs: list[dict]) -> list[tuple[int, str]]:
        """Try structured heading styles (Heading 1/2/3)."""
        boundaries: list[tuple[int, str]] = []
        for i, p in enumerate(paragraphs):
            if not p["is_heading"] or not p["text"]:
                continue
            text_lower = p["text"].lower()
            for section_name, keywords in self.SECTION_KEYWORDS.items():
                if any(kw in text_lower for kw in keywords):
                    boundaries.append((i, section_name))
                    break
        boundaries.sort()
        return boundaries

    def _segment_by_visual_cues(self, paragraphs: list[dict]) -> list[tuple[int, str]]:
        """Detect section headings by visual formatting.

        Catches paragraphs that are:
        - Short (<= 80 chars)
        - AND (all-caps OR (bold + size >= 14pt))
        """
        boundaries: list[tuple[int, str]] = []

        for i, p in enumerate(paragraphs):
            text = p["text"]
            if not text or len(text) > 80:
                continue

            is_all_caps = text == text.upper() and len(text) > 2
            is_bold_large = p["bold"] and (p["size"] or 0) >= 14

            if not (is_all_caps or is_bold_large):
                continue

            text_lower = text.lower().rstrip(".:：。.")
            for section_name, keywords in self.SECTION_KEYWORDS.items():
                if any(kw == text_lower or kw in text_lower for kw in keywords):
                    boundaries.append((i, section_name))
                    break

        boundaries.sort()
        return boundaries

    def _apply_sections(self, paper: ParsedPaper, boundaries: list[tuple[int, str]]) -> None:
        """Map paragraphs to ParsedPaper section fields using boundary indices."""
        # Reconstruct full text of all paragraphs for slicing
        all_text = paper.full_text.split("\n\n")

        for idx, (start_idx, name) in enumerate(boundaries):
            end_idx = boundaries[idx + 1][0] if idx + 1 < len(boundaries) else len(all_text)
            section_paras = all_text[start_idx:end_idx]
            section_text = "\n\n".join(section_paras).strip()
            if section_text:
                setattr(paper, name, section_text)

        # If abstract not found via heading, try first substantial text
        if not paper.abstract:
            self._fallback_abstract(paper)

    def _fallback_abstract(self, paper: ParsedPaper) -> None:
        """Use first 500 chars as abstract if no structured abstract found."""
        text = paper.full_text
        if not text:
            return
        for para in text.split("\n\n"):
            stripped = para.strip()
            if len(stripped) > 150:
                paper.abstract = stripped[:3000]
                return
        paper.abstract = text[:2000]

    def _extract_metadata(self, paper: ParsedPaper, doc) -> None:
        """Extract title, authors from core properties."""
        try:
            props = doc.core_properties
            if props.title:
                paper.title = props.title
            if props.author:
                paper.authors = [a.strip() for a in props.author.split(";")]
        except Exception:
            pass

    def _extract_tables(self, paper: ParsedPaper, doc) -> None:
        """Extract tables into list[dict] with headers as keys."""
        for table in doc.tables:
            rows = table.rows
            if len(rows) < 2:
                continue
            headers = [cell.text.strip() for cell in rows[0].cells]
            if not any(headers):
                headers = [f"col_{i}" for i in range(len(headers))]
            data = []
            for row in rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    data.append(dict(zip(headers, cells)))
            if data:
                paper.tables.append({"headers": headers, "rows": data})

    def _extract_images(self, paper: ParsedPaper, doc, image_dir: Path) -> None:
        """Extract embedded images from docx relationships."""
        count = 0
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype:
                continue
            try:
                image = rel.target_part
                ext = Path(image.partname).suffix or ".png"
                if ext.lower() not in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif"):
                    ext = ".png"
            except Exception:
                continue

            count += 1
            dest = image_dir / f"docx_image_{count}{ext}"
            try:
                dest.write_bytes(image.blob)
                paper.image_paths.append(dest)
            except Exception:
                pass
